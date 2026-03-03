from __future__ import annotations

import csv
import glob
import json
import sys
import time
from pathlib import Path

import typer
import yaml
from haystack import Pipeline
from haystack.components.converters import (
    HTMLToDocument,
    MarkdownToDocument,
    PyPDFToDocument,
    TextFileToDocument,
)
from haystack.components.preprocessors import DocumentSplitter
from haystack.components.writers import DocumentWriter
from haystack.dataclasses import Document
from haystack.document_stores.types import DuplicatePolicy
from haystack_integrations.document_stores.qdrant import QdrantDocumentStore
from qdrant_client import QdrantClient
from qdrant_client.http.models import FieldCondition, Filter, MatchValue

from app.core.config import Settings, get_settings
from ingestion.pipelines.components import LlamaCppEmbedding

DEFAULT_CONFIG_PATH = Path("ingestion/config/filesystem.example.yaml")

app = typer.Typer(help="Ingestion CLI for filesystem-based documents")
INPUTS_ARG = typer.Argument(
    None,
    help="File paths or glob patterns to ingest (optional; overrides config sources).",
)


def load_yaml_config(path: Path) -> dict:
    if not path.exists():
        return {}

    with path.open("r", encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def normalize_document(doc: Document) -> Document:
    """Lightweight normalization to trim whitespace and ensure meta exists."""

    content = doc.content or ""
    normalized_lines = [line.strip() for line in content.splitlines()]
    doc.content = "\n".join(normalized_lines).strip()
    if doc.meta is None:
        doc.meta = {}
    return doc


def load_documents_from_paths(
    paths: list[str], base_dir: Path
) -> tuple[list[Document], list[Path]]:
    documents: list[Document] = []
    pdf_converter = PyPDFToDocument()
    md_converter = MarkdownToDocument()
    html_converter = HTMLToDocument()
    txt_converter = TextFileToDocument()

    matched_files: set[Path] = set()
    for pattern in paths:
        resolved_pattern = Path(pattern)
        candidate_patterns: list[Path]
        if resolved_pattern.is_absolute():
            candidate_patterns = [resolved_pattern]
        else:
            cwd = Path.cwd().resolve()
            candidate_patterns = [(base_dir / resolved_pattern).resolve()]
            if cwd != base_dir:
                candidate_patterns.append((cwd / resolved_pattern).resolve())

        for candidate in candidate_patterns:
            for file_path in glob.glob(str(candidate), recursive=True):
                matched_files.add(Path(file_path))

    if not matched_files:
        typer.echo(f"[ingest] No files matched patterns: {paths}", err=True)
        return documents, []

    typer.echo(f"[ingest] Matched {len(matched_files)} files for conversion")

    sorted_files = sorted(matched_files)
    for path in sorted_files:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            converter = pdf_converter
        elif suffix in {".md", ".markdown"}:
            converter = md_converter
        elif suffix in {".html", ".htm"}:
            converter = html_converter
        else:
            converter = txt_converter

        if suffix == ".docx":
            try:
                import docx as python_docx  # noqa: PLC0415

                word_doc = python_docx.Document(str(path))
                text = "\n".join(p.text for p in word_doc.paragraphs if p.text.strip())
                raw_docs = [Document(content=text, meta={})]
            except Exception as exc:  # pragma: no cover
                typer.echo(f"[ingest] Failed to convert DOCX {path}: {exc}", err=True)
                continue
            for doc in raw_docs:
                if doc.meta is None:
                    doc.meta = {}
                doc.meta["source"] = str(path)
                documents.append(normalize_document(doc))
            continue

        if suffix == ".csv":
            try:
                with path.open(newline="", encoding="utf-8-sig") as fh:
                    reader = csv.reader(fh)
                    rows = [", ".join(row) for row in reader]
                text = "\n".join(rows)
                raw_docs = [Document(content=text, meta={})]
            except Exception as exc:  # pragma: no cover
                typer.echo(f"[ingest] Failed to convert CSV {path}: {exc}", err=True)
                continue
            for doc in raw_docs:
                if doc.meta is None:
                    doc.meta = {}
                doc.meta["source"] = str(path)
                documents.append(normalize_document(doc))
            continue

        try:
            result = converter.run(sources=[path])
        except Exception as exc:  # pragma: no cover - best-effort logging
            typer.echo(f"[ingest] Failed to convert {path}: {exc}", err=True)
            continue

        for doc in result.get("documents", []):
            if doc.meta is None:
                doc.meta = {}
            doc.meta["source"] = str(path)
            documents.append(normalize_document(doc))

    return documents, sorted_files


def _group_documents_by_source(docs: list[Document]) -> dict[str, list[Document]]:
    grouped: dict[str, list[Document]] = {}
    for doc in docs:
        meta = doc.meta or {}
        source_val = meta.get("source") or meta.get("file_path") or "unknown"
        grouped.setdefault(str(source_val), []).append(doc)
    return grouped


def _split_large_documents(docs: list[Document], max_chars: int) -> list[Document]:
    if max_chars <= 0:
        return docs
    output: list[Document] = []
    for doc in docs:
        content = doc.content or ""
        if len(content) <= max_chars:
            output.append(doc)
            continue
        meta_base = dict(doc.meta or {})
        parts = _split_text_with_separators(content, max_chars)
        for idx, chunk in enumerate(parts, start=1):
            meta = dict(meta_base)
            meta["source_part"] = idx
            output.append(Document(content=chunk, meta=meta))
    return output


def _split_text_with_separators(text: str, max_chars: int) -> list[str]:
    if max_chars <= 0 or len(text) <= max_chars:
        return [text]
    separators = ["\n\n", "\n", " "]
    segments = [text]
    for sep in separators:
        next_segments: list[str] = []
        for segment in segments:
            if len(segment) <= max_chars:
                next_segments.append(segment)
                continue
            if sep not in segment:
                next_segments.append(segment)
                continue
            current = ""
            for chunk in segment.split(sep):
                if not chunk:
                    continue
                candidate = f"{current}{sep}{chunk}" if current else chunk
                if len(candidate) <= max_chars:
                    current = candidate
                else:
                    if current:
                        next_segments.append(current)
                    if len(chunk) > max_chars:
                        next_segments.append(chunk)
                        current = ""
                    else:
                        current = chunk
            if current:
                next_segments.append(current)
        segments = next_segments
    flattened: list[str] = []
    for segment in segments:
        if len(segment) <= max_chars:
            flattened.append(segment.strip())
            continue
        for start in range(0, len(segment), max_chars):
            flattened.append(segment[start : start + max_chars].strip())
    return [seg for seg in flattened if seg]


def _source_value_candidates(path: Path) -> list[str]:
    resolved = path.resolve()
    candidates = {
        str(path),
        str(resolved),
        str(path).replace("\\", "/"),
        str(resolved).replace("\\", "/"),
        path.as_posix(),
        resolved.as_posix(),
        path.name,
    }
    return sorted({val for val in candidates if val})


def purge_existing_documents(settings: Settings, collection: str, sources: list[Path]) -> None:
    if not sources:
        return
    client = QdrantClient(
        host=settings.qdrant_host,
        port=settings.qdrant_port,
        api_key=settings.qdrant_api_key,
        timeout=30.0,
    )
    for source_path in sources:
        source_vals = _source_value_candidates(source_path)
        keys = ("source", "file_path", "meta.source", "meta.file_path")
        conditions = [
            FieldCondition(key=key, match=MatchValue(value=val))
            for key in keys
            for val in source_vals
        ]
        flt = Filter(should=conditions)
        try:
            count_result = client.count(collection_name=collection, count_filter=flt, exact=True)
            count_val = getattr(count_result, "count", None)
            typer.echo(
                f"[ingest] Purging existing docs collection={collection} "
                f"source={source_path} match_count={count_val}"
            )
            typer.echo(f"[ingest] Purge match keys={list(keys)} values={source_vals}")
            result = client.delete(collection_name=collection, points_selector=flt, wait=True)
            status = getattr(result, "status", None)
            if status is not None:
                typer.echo(
                    f"[ingest] Purge result collection={collection} source={source_path} "
                    f"status={status} deleted_count={count_val}"
                )
        except Exception as exc:  # pragma: no cover - best-effort cleanup
            typer.echo(
                f"[ingest] Failed to purge existing docs for {source_path}: {exc}",
                err=True,
            )


def build_filesystem_ingestion_pipeline(
    settings: Settings,
    collection: str = "documents",
    vector_size: int = 1024,
    split_by: str = "word",
    split_length: int = 128,
    split_overlap: int = 32,
) -> Pipeline:
    document_store = QdrantDocumentStore(
        url=f"http://{settings.qdrant_host}:{settings.qdrant_port}",
        api_key=settings.qdrant_api_key,
        embedding_dim=vector_size,
        index=collection,
        prefer_grpc=False,
    )

    pipeline = Pipeline()
    pipeline.add_component(
        "splitter",
        DocumentSplitter(
            split_by=split_by,
            split_length=split_length,
            split_overlap=split_overlap,
        ),
    )
    pipeline.add_component(
        "embedder",
        LlamaCppEmbedding(endpoint=str(settings.llama_embedding_url), model="bge-large-en-v1.5"),
    )
    pipeline.add_component(
        "writer",
        DocumentWriter(document_store=document_store, policy=DuplicatePolicy.OVERWRITE),
    )

    pipeline.connect("splitter", "embedder")
    pipeline.connect("embedder", "writer")
    return pipeline


@app.command()
def ingest(
    inputs: list[str] = INPUTS_ARG,
    paths: str | None = typer.Option(
        None,
        "--path",
        "-p",
        help="Glob pattern for PDFs/HTML/Markdown to ingest (overrides config sources)",
        show_default="from config",
    ),
    config: Path = typer.Option(  # noqa: B008
        DEFAULT_CONFIG_PATH,
        "--config",
        "-c",
        help="Path to YAML config defining sources/chunking settings",
        exists=True,
        dir_okay=False,
        readable=True,
        resolve_path=True,
    ),
    base_dir: Path | None = typer.Option(  # noqa: B008
        None,
        "--base-dir",
        "-b",
        help="Base directory for resolving relative glob patterns (defaults to config file dir)",
        dir_okay=True,
        file_okay=False,
        resolve_path=True,
    ),
    collection: str | None = typer.Option(  # noqa: B008
        None,
        "--collection",
        "-n",
        help="Qdrant collection name (overrides config and --tenant-id when provided)",
    ),
    tenant_id: str | None = typer.Option(  # noqa: B008
        None,
        "--tenant-id",
        help=(
            "Tenant slug; resolves collection as 'tenant_<id>'"
            " (required unless --collection is set)"
        ),
    ),
    artifacts_dir: Path | None = typer.Option(  # noqa: B008
        None,
        "--artifacts-dir",
        help="Optional directory to dump cleaned documents before splitting",
        dir_okay=True,
        file_okay=False,
        resolve_path=True,
    ),
):
    settings = get_settings()
    if collection is None and tenant_id is None:
        typer.echo("[ingest] Error: provide --tenant-id <slug> or --collection <name>", err=True)
        raise typer.Exit(code=1)
    typer.echo(f"[ingest] Loading config from {config}")
    config_data = load_yaml_config(config)

    config_base_dir = Path(config_data.get("base_dir")) if config_data.get("base_dir") else None
    if (inputs or paths) and base_dir is None:
        active_base_dir = Path.cwd().resolve()
    else:
        active_base_dir = (base_dir or config_base_dir or config.parent).resolve()

    source_globs: list[str] = []
    if inputs:
        source_globs.extend(inputs)
    if paths:
        source_globs.append(paths)
    if not source_globs:
        source_globs = config_data.get("sources", [])
    if collection:
        selected_collection = collection
    elif tenant_id:
        selected_collection = f"tenant_{tenant_id}"
    else:
        selected_collection = config_data.get("collection", "documents")
    chunking = config_data.get("chunking", {})
    split_by = chunking.get("split_by", "word")
    split_length = int(chunking.get("split_length", 200))
    split_overlap = int(chunking.get("split_overlap", 40))
    max_document_chars = int(config_data.get("max_document_chars", 8000))

    if not source_globs:
        typer.echo("[ingest] No sources provided via CLI or config; nothing to ingest", err=True)
        raise typer.Exit(code=1)

    typer.echo(f"[ingest] Using base_dir={active_base_dir}")
    typer.echo(f"[ingest] Using sources: {source_globs}")
    docs, matched_files = load_documents_from_paths(source_globs, base_dir=active_base_dir)
    typer.echo(f"[ingest] Loaded {len(docs)} documents before chunking")
    if not docs:
        typer.echo("[ingest] No documents found for provided paths", err=True)
        raise typer.Exit(code=1)

    purge_existing_documents(settings, selected_collection, matched_files)

    typer.echo(
        "[ingest] Chunking with "
        f"split_by='{split_by}', length={split_length}, overlap={split_overlap}"
    )
    pipeline = build_filesystem_ingestion_pipeline(
        settings=settings,
        collection=selected_collection,
        split_by=split_by,
        split_length=split_length,
        split_overlap=split_overlap,
    )
    typer.echo(f"[ingest] Running pipeline into collection '{selected_collection}' per file")
    grouped_docs = _group_documents_by_source(docs)
    total_written = 0
    for source, source_docs in grouped_docs.items():
        prepared_docs = _split_large_documents(source_docs, max_document_chars)
        typer.echo(f"[ingest] Running pipeline for source={source} docs={len(prepared_docs)}")
        result = pipeline.run({"splitter": {"documents": prepared_docs}})
        writer_stats = result.get("writer", {})
        written = writer_stats.get("documents_written") or writer_stats.get("count") or 0
        total_written += int(written)
        typer.echo(f"[ingest] Pipeline finished source={source} wrote={written} chunks")
    typer.echo(f"[ingest] Pipeline finished; wrote {total_written} chunks")

    artifact_root = artifacts_dir or config_data.get("artifacts_dir")
    if artifact_root:
        artifact_path = Path(artifact_root)
        artifact_path.mkdir(parents=True, exist_ok=True)
        output_file = artifact_path / f"ingestion_{int(time.time())}.ndjson"
        with output_file.open("w", encoding="utf-8") as f:
            for doc in docs:
                json.dump(
                    {
                        "id": doc.id,
                        "content": doc.content,
                        "meta": doc.meta or {},
                    },
                    f,
                    ensure_ascii=False,
                )
                f.write("\n")
        typer.echo(f"[ingest] Dumped cleaned documents to {output_file}")

    typer.echo(
        f"[ingest] Completed ingestion into '{selected_collection}' with "
        f"{total_written} total chunks"
    )


if __name__ == "__main__":
    # Allow invoking with an optional "ingest" alias: `python -m ingestion.cli ingest ...`
    if len(sys.argv) > 1 and sys.argv[1] == "ingest":
        sys.argv.pop(1)
    app()
